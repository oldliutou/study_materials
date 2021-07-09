# 读取docx中的文本代码示例
import docx

# 获取文档对象
file = docx.Document("C:\\Users\\Administrator\\Desktop\\2019论文题目李鑫.docx")
# print("段落数:" + str(len(file.paragraphs)))  # 段落数为13，每个回车隔离一段

# 输出每一段的内容
# for para in file.paragraphs:
#     print(para.text)
#
# with open('fileContent.txt', 'a') as f:
#     f.write(para.text)
#     f.close()
# 输出段落编号及段落内容
for i in range(len(file.paragraphs)):
    # print("第" + str(i) + "段的内容是：" + file.paragraphs[i].text)
    print('['+file.paragraphs[i].text[:4]+']'+file.paragraphs[i].text[4:])

    # 2019论文题目  李鑫
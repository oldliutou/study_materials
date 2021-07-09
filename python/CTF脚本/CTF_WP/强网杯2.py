import os
import docx
path = 'C:\\Users\\Administrator\\Desktop\\five_month'
def get_filelist(dir):
    Filelist = []
    for home, dirs, files in os.walk(path):
        for filename in files:
 # 文件名列表，包含完整路径
            Filelist.append(os.path.join(home, filename))
 # # 文件名列表，只包含文件名
# Filelist.append( filename)
    return Filelist
if __name__ == "__main__":
    Filelist = get_filelist(dir)
    print(len(Filelist))

    for file in Filelist:
        try:
            print(file)
            file = docx.Document(file)

            # 输出每一段的内容
            for para in file.paragraphs:
                print(para.text)
            # 输出段落编号及段落内容
            # for i in range(len(file.paragraphs)):
            #     filecon = file+"第" + str(i) + "段的内容是：" + file.paragraphs[i].text
            #     print("第" + str(i) + "段的内容是：" + file.paragraphs[i].text)
                with open('fileContent.txt','a+') as f:
                    f.write(para.text)
                    f.close()
        except:
            pass

